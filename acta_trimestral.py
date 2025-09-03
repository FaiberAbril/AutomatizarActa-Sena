import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import os
import shutil
from docx import Document
import tempfile

def acta_trimestral():
    st.header("Generación de Acta Trimestral")
    
    # Cargar archivos
    st.subheader("1. Cargar archivos necesarios")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Reporte de Juicios Evaluativos**")
        archivo_juicios = st.file_uploader("Seleccione el archivo Excel", 
                                        type=['xls', 'xlsx'],
                                        key="juicios")
    
    with col2:
        st.markdown("**Reporte por instructor**")
        archivo_instructor = st.file_uploader("Seleccione el archivo Excel", 
                                            type=['xls', 'xlsx'],
                                            key="instructor")
    
    # Parámetros del acta
    st.subheader("2. Configuración del acta")
    
    col1, col2 = st.columns(2)
    
    with col1:
        programa = st.text_input("Nombre del Programa de Formación")
        trimestre = st.selectbox("Trimestre", ["I", "II", "III", "IV"])
    
    with col2:
        instructor = st.text_input("Nombre del Instructor")
        fecha = st.date_input("Fecha del acta")

    # Botón para generar el acta
    if st.button("Generar Acta", type="primary"):
        if not all([archivo_juicios, archivo_instructor]):
            st.error("Por favor, cargue todos los archivos necesarios.")
            return
        
        try:
            with st.spinner("Generando acta, por favor espere..."):
                # Guardar archivos temporalmente
                temp_dir = tempfile.mkdtemp()
                
                ruta_juicios = os.path.join(temp_dir, "Reporte de Juicios Evaluativos.xls")
                with open(ruta_juicios, "wb") as f:
                    f.write(archivo_juicios.getbuffer())
                
                ruta_instructor = os.path.join(temp_dir, "Reporteporinstructor.xls")
                with open(ruta_instructor, "wb") as f:
                    f.write(archivo_instructor.getbuffer())
                
                # Ruta de la plantilla que ya está en el proyecto
                ruta_plantilla_origen = "acta.docx"  # Ajusta esta ruta según donde esté tu plantilla
                
                # Copiar la plantilla al directorio temporal
                ruta_plantilla = os.path.join(temp_dir, "acta.docx")
                shutil.copy(ruta_plantilla_origen, ruta_plantilla)
                
                # Generar el acta
                ruta_destino = os.path.join(temp_dir, "actagenerada.docx")
                generar_acta_completa(ruta_juicios, ruta_instructor, ruta_plantilla, ruta_destino,
                                    programa, trimestre, fecha, instructor)
                
                # Descargar el archivo generado
                with open(ruta_destino, "rb") as f:
                    bytes_data = f.read()
                
                st.success("¡Acta generada con éxito!")
                st.download_button(
                    label="Descargar acta generada",
                    data=bytes_data,
                    file_name=f"acta_trimestral_{programa.replace(' ', '_')}_{trimestre}_{fecha}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Limpiar archivos temporales
                shutil.rmtree(temp_dir)
                
        except Exception as e:
            st.error(f"Error al generar el acta: {str(e)}")

def generar_acta_completa(ruta_juicios, ruta_instructor, ruta_plantilla, ruta_destino,
                    programa, trimestre, fecha, instructor):
    
    """Función para generar el acta completa con los datos proporcionados"""
    
    # Copiar la plantilla
    shutil.copy(ruta_plantilla, ruta_destino)
    
    # Abrir el documento Word
    documento_word = Document(ruta_destino)
    
    # Cargar los archivos de Excel
    df = pd.read_excel(ruta_juicios, header=0, skiprows=12)
    dfinstructorhoras = pd.read_excel(ruta_instructor, header=0, skiprows=10)
    
    # Acceder a la tabla específica del documento
    target_table = documento_word.tables[0]
    target_cell = target_table.cell(7, 1)

    # Asegúrate de tener estas variables en tu función o parámetros
    ciudad = "Bogotá"  # o obtén esta información del usuario
    lugar = "Presencial"  # o "Enlace virtual: https://..."
    centro_formacion = "Centro de Formación Norte"  # o el valor correspondiente
    fecha = datetime.now()  # o usa la fecha que ya tienes



    target_table.cell(1, 0).text = f"NOMBRE DEL COMITÉ O DE LA REUNIÓN: SEGUIMIENTO FORMATIVO - {programa.upper()} - FICHA xxxxx - TRIMESTRE {trimestre}"

    # Fila 2, Columna 0: Ciudad y Fecha
    target_table.cell(2, 0).text = f"CIUDAD Y FECHA: {ciudad}, {fecha.strftime('%d/%m/%Y')}"

    # Fila 2, Columna 1: Hora Inicio
    target_table.cell(2, 1).text = "HORA INICIO: 08:00 AM"

    # Fila 2, Columna 2: Hora Fin  
    target_table.cell(2, 2).text = "HORA FIN: 05:00 PM"

    # Fila 3, Columna 0: Lugar y/o Enlace
    target_table.cell(3, 2).text = f"LUGAR Y/O ENLACE: {lugar}"

    # Fila 3, Columna 1: Dirección / Regional / Centro
    target_table.cell(3, 3).text = f"DIRECCIÓN / REGIONAL / CENTRO: {centro_formacion}"

    
    # Crear una nueva columna 'Nombre completo' concatenando 'Nombre' y 'Apellidos'
    df['Nombrecompleto'] = df['Nombre'] + ' ' + df['Apellidos']
    
    # Información para la tabla de estudiantes
    dfestudiantes = df.loc[df['Estado'] == 'EN FORMACION', ['Nombrecompleto', 'Número de Documento', 'Estado']]
    dfestudiantes = dfestudiantes.drop_duplicates(subset=['Nombrecompleto'])
    
    # Buscar la frase dentro de la celda y agregar tabla
    for paragraph in target_cell.paragraphs:
        if "Lo cual indica que se encuentran" in paragraph.text:
            # Agregar una tabla al final del documento
            tabla = documento_word.add_table(rows=len(dfestudiantes) + 1, cols=3)
            tabla.style = 'Table Grid'
            
            # Encabezado de la tabla
            encabezado = tabla.rows[0].cells
            encabezado[0].text = 'Nombre Completo'
            encabezado[1].text = 'Número de Documento'
            encabezado[2].text = 'Estado'
            
            # Datos de la tabla
            for i, (_, fila) in enumerate(dfestudiantes.iterrows(), start=1):
                celdas = tabla.rows[i].cells
                celdas[0].text = fila['Nombrecompleto']
                celdas[1].text = str(fila['Número de Documento'])
                celdas[2].text = fila['Estado']
            
            # Insertar la tabla
            paragraph._p.addnext(tabla._element)
            break
    
    # Agregar un párrafo en blanco
    documento_word.add_paragraph()
    

    # Buscar la frase dentro de la celda y agregar tabla
    for paragraph in target_cell.paragraphs:
        if "Verificación de asistencia a la reunión " in paragraph.text:
            # Agregar el nombre del instructor después de la frase
            nuevo_parrafo = documento_word.add_paragraph(f"Se confirma la asistencia y participación del instructor asignado como líder de ficha {instructor}, de acuerdo con la lista de asistencia anexa.")
            
            # Insertar el nuevo párrafo después del párrafo encontrado
            paragraph._p.addnext(nuevo_parrafo._element)
            break

    # Agregar un párrafo en blanco
    documento_word.add_paragraph()







    # Información para la tabla de estudiantes por retiros
    dfestudiantesretiro = df.loc[df['Estado'] != 'EN FORMACION', ['Nombrecompleto', 'Número de Documento', 'Estado']]
    dfestudiantesretiro = dfestudiantesretiro.drop_duplicates(subset=['Nombrecompleto'])
    
    # Buscar la frase para la tabla de retiros
    for paragraph in target_cell.paragraphs:
        if "tabla de los cancelados y novedades de retiro se puede eliminar de la tabla anterior" in paragraph.text:
            # Agregar tabla de retiros
            tablaRetiro = documento_word.add_table(rows=len(dfestudiantesretiro) + 1, cols=3)
            tablaRetiro.style = 'Table Grid'
            
            # Encabezado de la tabla
            encabezado = tablaRetiro.rows[0].cells
            encabezado[0].text = 'Nombre Completo'
            encabezado[1].text = 'Número de Documento'
            encabezado[2].text = 'Estado'
            
            # Datos de la tabla
            for i, (_, fila) in enumerate(dfestudiantesretiro.iterrows(), start=1):
                celdas = tablaRetiro.rows[i].cells
                celdas[0].text = fila['Nombrecompleto']
                celdas[1].text = str(fila['Número de Documento'])
                celdas[2].text = fila['Estado']
            
            # Insertar la tabla
            paragraph._p.addnext(tablaRetiro._element)
            break
    
    # Agregar un párrafo en blanco entre las tablas
    documento_word.add_paragraph()

    # Cargar el archivo de Excel en competencias
    dfcompetencias = df
    dfcompetencias = dfcompetencias[dfcompetencias['Competencia'] != '2 - RESULTADOS DE APRENDIZAJE ETAPA PRACTICA']
    dfcompetencias['Juicio de Evaluación'] = dfcompetencias['Juicio de Evaluación'].apply(
        lambda x: 'Si' if x == 'APROBADO' else 'No' if x == 'POR EVALUAR' else x)
    dfcompetencias = dfcompetencias.drop_duplicates(subset=['Competencia'])

    # Buscar la frase dentro de la celda
    for paragraph in target_cell.paragraphs:
        if "3.2.1 Competencias Evaluadas" in paragraph.text:
            # Agregar una tabla al final del documento
            tablaCompetencias = documento_word.add_table(rows=len(dfcompetencias) + 1, cols=2)
            tablaCompetencias.style = 'Table Grid'

            # Agregar el encabezado de la tabla
            encabezado = tablaCompetencias.rows[0].cells
            encabezado[0].text = 'Competencias Desarrolladas'
            encabezado[1].text = 'Evaluado En SOFIAPLUS'

            # Agregar los datos de las columnas especificadas a la tabla
            for i, (_, fila) in enumerate(dfcompetencias.iterrows(), start=1):
                celdas = tablaCompetencias.rows[i].cells
                celdas[0].text = fila['Competencia']
                celdas[1].text = str(fila['Juicio de Evaluación'])

            # Insertar la tabla debajo de la frase encontrada
            paragraph._p.addnext(tablaCompetencias._element)
            break

    # Agregar un párrafo en blanco entre las tablas
    documento_word.add_paragraph()

    # Resultados y competencias sin evaluar
    dfcompetenciassinevaluar = df
    dfcompetenciassinevaluar = dfcompetenciassinevaluar[['Competencia', 'Resultado de Aprendizaje', 'Juicio de Evaluación']]
    dfcompetenciassinevaluar = dfcompetenciassinevaluar[
        dfcompetenciassinevaluar['Competencia'] != '2 - RESULTADOS DE APRENDIZAJE ETAPA PRACTICA']
    dfcompetenciassinevaluar['Juicio de Evaluación'] = dfcompetenciassinevaluar['Juicio de Evaluación'].apply(
        lambda x: 'Si' if x == 'APROBADO' else 'No' if x == 'POR EVALUAR' else x)

    # Filtrar las filas donde el conteo de 'Sí' es igual a cero y obtener las columnas 'Competencia' y 'Resultado de Aprendizaje'
    dfcompetenciassinevaluar = dfcompetenciassinevaluar.groupby('Resultado de Aprendizaje').agg(
        {'Juicio de Evaluación': lambda x: (x == 'Si').sum(), 'Competencia': 'first'}).reset_index()
    dfcompetenciassinevaluar = dfcompetenciassinevaluar[dfcompetenciassinevaluar['Juicio de Evaluación'] == 0][
        ['Competencia', 'Resultado de Aprendizaje']]

    # Buscar la frase dentro de la celda
    for paragraph in target_cell.paragraphs:
        if "3.2.2 Competencias y Resultados de Aprendizaje por Evaluar" in paragraph.text:
            # Agregar una tabla al final del documento
            tablaCompetenciassinevaluar = documento_word.add_table(rows=len(dfcompetenciassinevaluar) + 1, cols=2)
            tablaCompetenciassinevaluar.style = 'Table Grid'

            # Agregar el encabezado de la tabla
            encabezado = tablaCompetenciassinevaluar.rows[0].cells
            encabezado[0].text = 'Competencia'
            encabezado[1].text = 'Resultado de Aprendizaje sin Evaluar'

            # Iterar sobre las filas de la tabla y los datos del DataFrame
            for i, (index, fila) in enumerate(dfcompetenciassinevaluar.iterrows(), start=1):
                # Acceder a las celdas de la fila
                celdas = tablaCompetenciassinevaluar.rows[i].cells
                # Asignar los datos a las celdas
                celdas[0].text = str(fila['Competencia'])
                celdas[1].text = str(fila['Resultado de Aprendizaje'])

            # Insertar la tabla debajo de la frase encontrada
            paragraph._p.addnext(tablaCompetenciassinevaluar._element)
            break

    # Agregar un párrafo en blanco entre las tablas
    documento_word.add_paragraph()

    # Horas de planeación
    horas_planeacion = {
        'Utilizar herramientas informáticas de acuerdo con las necesidades de manejo de información': 48,
        'APLICAR PRÁCTICAS DE PROTECCIÓN AMBIENTAL, SEGURIDAD Y SALUD EN EL TRABAJO DE ACUERDO CON LAS POLÍTICAS ORGANIZACIONALES Y LA NORMATIVIDAD VIGENTE.': 48,
        'Razonar cuantitativamente frente a situaciones susceptibles de ser abordadas de manera matemática en contextos laborales, sociales y personales.': 48,
        'APLICACIÓN DE CONOCIMIENTOS DE LAS CIENCIAS NATURALES DE ACUERDO CON SITUACIONES DEL CONTEXTO PRODUCTIVO Y SOCIAL.': 48,
        'DESARROLLAR PROCESOS DE COMUNICACIÓN EFICACES Y EFECTIVOS, TENIENDO EN CUENTA SITUACIONES  DE ORDEN SOCIAL, PERSONAL Y PRODUCTIVO.': 48,
        'GENERAR HÁBITOS SALUDABLES DE VIDA MEDIANTE LA APLICACIÓN DE PROGRAMAS DE ACTIVIDAD FÍSICA EN LOS CONTEXTOS PRODUCTIVOS Y SOCIALES.': 48,
        'Orientar investigación formativa según referentes técnicos': 48,
        'Enrique Low Murtra-Interactuar en el contexto productivo y social de acuerdo con principios  éticos para la construcción de una cultura de paz.': 48,
        'INTERACTUAR EN LENGUA INGLESA DE FORMA ORAL Y ESCRITA DENTRO DE CONTEXTOS SOCIALES Y LABORALES SEGÚN LOS CRITERIOS ESTABLECIDOS POR EL MARCO COMÚN EUROPEO DE REFERENCIA PARA LAS LENGUAS.': 384,
        'Fomentar cultura emprendedora según habilidades y competencias personales': 48
    }

    dfinstructorhoras['Nombrecompleto'] = dfinstructorhoras['Nombre Instructor'] + ' ' + dfinstructorhoras['Apellido Instructor']
    dfinstructorhoras = dfinstructorhoras.sort_values(by='Competencia', ascending=True)

    # Añadir la columna 'Horas Planeadas' al DataFrame
    dfinstructorhoras['Horas Planeadas'] = dfinstructorhoras['Competencia'].map(horas_planeacion)

    dfinstructorhoras = dfinstructorhoras[['Nombrecompleto', 'Estado Instructor', 'Competencia', 'Horas Programadas', 'Horas Planeadas']]

    # Buscar la frase dentro de la celda
    for paragraph in target_cell.paragraphs:
        if "El reporte de instructor por ficha en SOFIA PLUS es:" in paragraph.text:
            # Agregar una tabla al final del documento
            tablainstructores = documento_word.add_table(rows=len(dfinstructorhoras) + 1, cols=5)
            tablainstructores.style = 'Table Grid'

            # Agregar el encabezado de la tabla
            encabezado = tablainstructores.rows[0].cells
            encabezado[0].text = 'Nombre completo Instructor'
            encabezado[1].text = 'Estado Instructor'
            encabezado[2].text = 'Competencia'
            encabezado[3].text = 'Horas Programadas en sofia plus'
            encabezado[4].text = 'Horas Planeacion'

            # Iterar sobre las filas de la tabla y los datos del DataFrame
            for i, (index, fila) in enumerate(dfinstructorhoras.iterrows(), start=1):
                # Acceder a las celdas de la fila
                celdas = tablainstructores.rows[i].cells
                # Asignar los datos a las celdas
                celdas[0].text = str(fila['Nombrecompleto'])
                celdas[1].text = str(fila['Estado Instructor'])
                celdas[2].text = str(fila['Competencia'])
                celdas[3].text = str(fila['Horas Programadas'])
                celdas[4].text = str(fila['Horas Planeadas'])

            # Insertar la tabla debajo de la frase encontrada
            paragraph._p.addnext(tablainstructores._element)
            break

    # Agregar un párrafo en blanco entre las tablas
    documento_word.add_paragraph()

    # Guardar el documento generado
    documento_word.save(ruta_destino)

# Si se ejecuta este archivo directamente
if __name__ == "__main__":
    acta_trimestral()