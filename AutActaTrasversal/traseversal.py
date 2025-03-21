import pandas as pd
from docx import Document


# Abrir el documento Word existente
documento_word = Document('AutActaTrasversal/ACTATRASVERSAL.docx')

# Cargar el archivo de Excel en un DataFrame con la cabecera en la fila 13 y comenzando desde la fila 14
df = pd.read_excel('AutActaTrasversal/Reporte de Juicios Evaluativos.xls',header=0,skiprows=12)

dfcompetencia = pd.read_excel('AutActaTrasversal/Reporte de Juicios Evaluativos.xls',header=0,skiprows=12)

nombreCompetencia = "37371 - Utilizar herramientas informáticas de acuerdo con las necesidades de manejo de información"


# Access the specific paragraph where you want to insert text
target_table = documento_word.tables[0]  # replace 0 with the index of the table
target_cell = target_table.cell(7, 1)  # Suponiendo que la celda está en la fila 7 y columna 1
target_paragraph = None


# Crear una nueva columna 'Nombre completo' concatenando 'Nombre' y 'Apellidos'
df['Nombrecompleto'] = df['Nombre'] + ' ' + df['Apellidos']


# informacion para la tabla de estudiantes 
dfestudiantes=df.loc[df['Estado'] == 'EN FORMACION', ['Nombrecompleto', 'Número de Documento','Estado']]
dfestudiantes = dfestudiantes.drop_duplicates(subset=['Nombrecompleto'])


# Buscar la frase dentro de la celda
for paragraph in target_cell.paragraphs:
    if "Lo cual indica que se encuentran en formación:" in paragraph.text:
        # Agregar una tabla al final del documento
        tabla = documento_word.add_table(rows=len(dfestudiantes) + 1, cols=3)  # +1 para incluir el encabezado
        tabla.style = 'Table Grid'  # Aplicar un estilo de tabla

        # Agregar el encabezado de la tabla
        encabezado = tabla.rows[0].cells
        encabezado[0].text = 'Nombre Completo'
        encabezado[1].text = 'Número de Documento'
        encabezado[2].text = 'Estado'

        # Agregar los datos de las columnas especificadas a la tabla
        for i, (_, fila) in enumerate(dfestudiantes.iterrows(), start=1):
            celdas = tabla.rows[i].cells
            celdas[0].text = fila['Nombrecompleto']
            celdas[1].text = str(fila['Número de Documento'])  # Convertir a cadena si no es texto
            celdas[2].text = fila['Estado']

    # Insertar la tabla debajo de la frase encontrada
        paragraph._p.addnext(tabla._element)
        break


# Agregar un párrafo en blanco entre las tablas
documento_word.add_paragraph()


# informacion para la tabla de estudiantes por retiros 
dfestudiantes=df.loc[df['Estado'] != 'EN FORMACION', ['Nombrecompleto', 'Número de Documento','Estado']]
dfestudiantesretiro = dfestudiantes.drop_duplicates(subset=['Nombrecompleto'])


    # Buscar la frase dentro de la celda
for paragraph in target_cell.paragraphs:
    if "tabla de los cancelados y novedades de retiro se puede eliminar de la tabla anterior" in paragraph.text:
    
        # Agregar una tabla al final del documento
        tablaRetiro = documento_word.add_table(rows=len(dfestudiantesretiro) + 1, cols=3)  # +1 para incluir el encabezado
        tablaRetiro.style = 'Table Grid'  # Aplicar un estilo de tabla

        # Agregar el encabezado de la tabla
        encabezado = tablaRetiro.rows[0].cells
        encabezado[0].text = 'Nombre Completo'
        encabezado[1].text = 'Número de Documento'
        encabezado[2].text = 'Estado'

        # Agregar los datos de las columnas especificadas a la tabla
        for i, (_, fila) in enumerate(dfestudiantesretiro.iterrows(), start=1):
            celdas = tablaRetiro.rows[i].cells
            celdas[0].text = fila['Nombrecompleto']
            celdas[1].text = str(fila['Número de Documento'])  # Convertir a cadena si no es texto
            celdas[2].text = fila['Estado']

    # Insertar la tabla debajo de la frase encontrada
        paragraph._p.addnext(tablaRetiro._element)
        break


# Agregar un párrafo en blanco entre las tablas
documento_word.add_paragraph()



dfcompetencia = dfcompetencia[df['Competencia'] == nombreCompetencia]
dfcompetencia['Nombrecompleto'] = dfcompetencia['Nombre'] + ' ' + dfcompetencia['Apellidos']
dfcompetencia = dfcompetencia[['Número de Documento', 'Nombrecompleto', 'Juicio de Evaluación']]
dfcompetencia = dfcompetencia.drop_duplicates(subset='Nombrecompleto')
print(dfcompetencia)

numero_filas = len(dfcompetencia)


# Buscar la frase dentro de la celda
for paragraph in target_cell.paragraphs:
    if "Aprendices Aprobados:" in paragraph.text:
        # Agregar una tabla al final del documento
        tabla = documento_word.add_table(rows=numero_filas + 1, cols=3) # +1 para incluir el encabezado
        tabla.style = 'Table Grid'  # Aplicar un estilo de tabla

        # Agregar el encabezado de la tabla
        encabezado = tabla.rows[0].cells
        encabezado[0].text = 'Nombre Completo'
        encabezado[1].text = 'Número de Documento'
        encabezado[2].text = 'Juicio de Evaluación'

        # Agregar los datos de las columnas especificadas a la tabla
        for i, (_, fila) in enumerate(dfcompetencia.iterrows(), start=1):
            celdas = tabla.rows[i].cells
            celdas[0].text = fila['Nombrecompleto']
            celdas[1].text = str(fila['Número de Documento'])  # Convertir a cadena si no es texto
            celdas[2].text = fila['Juicio de Evaluación']

    # Insertar la tabla debajo de la frase encontrada
        paragraph._p.addnext(tabla._element)
        break


# Agregar un párrafo en blanco entre las tablas
documento_word.add_paragraph()

#Aprendices por Evaluar:


documento_word.save('AutActaTrasversal/ACTATRASVERSAL.docx')