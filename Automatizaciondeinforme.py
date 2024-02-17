import pandas as pd
from docx import Document


# Abrir el documento Word existente
documento_word = Document('acta.docx')

# Cargar el archivo de Excel en un DataFrame con la cabecera en la fila 13 y comenzando desde la fila 14
df = pd.read_excel('Reporte de Juicios Evaluativos (1).xls',header=0,skiprows=12)


# estudiantes en formacion 


# Access the specific paragraph where you want to insert text
target_table = documento_word.tables[0]  # replace 0 with the index of the table
target_cell = target_table.cell(7, 1)  # Suponiendo que la celda está en la fila 7 y columna 1
target_paragraph = None


# Crear una nueva columna 'Nombre completo' concatenando 'Nombre' y 'Apellidos'
df['Nombrecompleto'] = df['Nombre'] + ' ' + df['Apellidos']


# informacion para la tabla de estudiantes 
dfestudiantes=df.loc[df['Estado'] == 'EN FORMACION', ['Nombrecompleto', 'Número de Documento','Estado']]
dfestudiantes = dfestudiantes.drop_duplicates(subset=['Nombrecompleto'])


# Buscar la frase dentro de la celda
for paragraph in target_cell.paragraphs:
    if "Lo cual indica que se encuentran" in paragraph.text:
        # Agregar una tabla al final del documento
        tabla = documento_word.add_table(rows=len(dfestudiantes) + 1, cols=3)  # +1 para incluir el encabezado
        tabla.style = 'Table Grid'  # Aplicar un estilo de tabla

        # Agregar el encabezado de la tabla
        encabezado = tabla.rows[0].cells
        encabezado[0].text = 'Nombre Completo'
        encabezado[1].text = 'Número de Documento'
        encabezado[2].text = 'Estado'

        # Agregar los datos de las columnas especificadas a la tabla
        for i, (_, fila) in enumerate(dfestudiantes.iterrows(), start=1):
            celdas = tabla.rows[i].cells
            celdas[0].text = fila['Nombrecompleto']
            celdas[1].text = str(fila['Número de Documento'])  # Convertir a cadena si no es texto
            celdas[2].text = fila['Estado']

    # Insertar la tabla debajo de la frase encontrada
        paragraph._p.addnext(tabla._element)
        break





# Agregar un párrafo en blanco entre las tablas
documento_word.add_paragraph()



# informacion para la tabla de estudiantes por retiros 
dfestudiantes=df.loc[df['Estado'] != 'EN FORMACION', ['Nombrecompleto', 'Número de Documento','Estado']]
dfestudiantesretiro = dfestudiantes.drop_duplicates(subset=['Nombrecompleto'])





    # Buscar la frase dentro de la celda
for paragraph in target_cell.paragraphs:
    if "tabla de los cancelados y novedades de retiro se puede eliminar de la tabla anterior" in paragraph.text:
    
        # Agregar una tabla al final del documento
        tablaRetiro = documento_word.add_table(rows=len(dfestudiantesretiro) + 1, cols=3)  # +1 para incluir el encabezado
        tablaRetiro.style = 'Table Grid'  # Aplicar un estilo de tabla

        # Agregar el encabezado de la tabla
        encabezado = tablaRetiro.rows[0].cells
        encabezado[0].text = 'Nombre Completo'
        encabezado[1].text = 'Número de Documento'
        encabezado[2].text = 'Estado'

        # Agregar los datos de las columnas especificadas a la tabla
        for i, (_, fila) in enumerate(dfestudiantesretiro.iterrows(), start=1):
            celdas = tablaRetiro.rows[i].cells
            celdas[0].text = fila['Nombrecompleto']
            celdas[1].text = str(fila['Número de Documento'])  # Convertir a cadena si no es texto
            celdas[2].text = fila['Estado']

    # Insertar la tabla debajo de la frase encontrada
        paragraph._p.addnext(tablaRetiro._element)
        break


# Agregar un párrafo en blanco entre las tablas
documento_word.add_paragraph()


# Cargar el archivo de Excel en competencias
dfcompetencias = df
dfcompetencias = dfcompetencias[dfcompetencias['Competencia'] != '2 - RESULTADOS DE APRENDIZAJE ETAPA PRACTICA']
grupos_competencia = dfcompetencias.groupby('Competencia')[['Competencia', 'Juicio de Evaluación']]
dfcompetencias['Juicio de Evaluación'] = dfcompetencias['Juicio de Evaluación'].apply(lambda x: 'Si' if x == 'APROBADO' else 'No' if x == 'POR EVALUAR' else x)
dfcompetencias = dfcompetencias.drop_duplicates(subset=['Competencia'])



# Buscar la frase dentro de la celda
for paragraph in target_cell.paragraphs:
    if "3.2.1 Competencias Evaluadas" in paragraph.text:

        # Agregar una tabla al final del documento
        tablaCompetencias = documento_word.add_table(rows=len(dfcompetencias) + 1, cols=2)  # +1 para incluir el encabezado
        tablaCompetencias.style = 'Table Grid'  # Aplicar un estilo de tabla

        # Agregar el encabezado de la tabla
        encabezado = tablaCompetencias.rows[0].cells
        encabezado[0].text = 'Competencias Desarrolladas'
        encabezado[1].text = 'Evaluado En SOFIAPLUS'


        # Agregar los datos de las columnas especificadas a la tabla
        for i, (_, fila) in enumerate(dfcompetencias.iterrows(), start=1):
            celdas = tablaCompetencias.rows[i].cells
            celdas[0].text = fila['Competencia']
            celdas[1].text = str(fila['Juicio de Evaluación'])  # Convertir a cadena si no es texto

    # Insertar la tabla debajo de la frase encontrada
        paragraph._p.addnext(tablaCompetencias._element)
        break


# Agregar un párrafo en blanco entre las tablas
documento_word.add_paragraph()

# resultados y competencias sin evaluar
dfcompetenciassinevaluar = df
dfcompetenciassinevaluar = dfcompetenciassinevaluar[['Competencia','Resultado de Aprendizaje','Juicio de Evaluación']]
dfcompetenciassinevaluar = dfcompetenciassinevaluar[dfcompetenciassinevaluar['Competencia'] != '2 - RESULTADOS DE APRENDIZAJE ETAPA PRACTICA']
dfcompetenciassinevaluar['Juicio de Evaluación'] = dfcompetenciassinevaluar['Juicio de Evaluación'].apply(lambda x: 'Si' if x == 'APROBADO' else 'No' if x == 'POR EVALUAR' else x)


# Supongamos que dfcompetenciassinevaluar es tu DataFrame original
dfcompetenciassinevaluar = dfcompetenciassinevaluar.groupby('Resultado de Aprendizaje').agg({'Juicio de Evaluación': lambda x: (x == 'Si').sum(), 'Competencia': 'first'}).reset_index()

# Ahora, si deseas filtrar las filas donde el conteo de 'Sí' es igual a cero y obtener las columnas 'Competencia' y 'Resultado de Aprendizaje':
dfcompetenciassinevaluar = dfcompetenciassinevaluar[dfcompetenciassinevaluar['Juicio de Evaluación'] == 0][['Competencia', 'Resultado de Aprendizaje']]
# Obtener los datos de las competencias y los resultados de aprendizaje
dfcompetenciassinevaluar = dfcompetenciassinevaluar[['Competencia', 'Resultado de Aprendizaje']]



# Buscar la frase dentro de la celda
for paragraph in target_cell.paragraphs:
    if "3.2.2 Competencias y Resultados de Aprendizaje por Evaluar" in paragraph.text:          
        # Agregar una tabla al final del documento
        tablaCompetenciassinevaluar = documento_word.add_table(rows=len(dfcompetenciassinevaluar) + 1, cols=2)  # +1 para incluir el encabezado
        tablaCompetenciassinevaluar.style = 'Table Grid'  # Aplicar un estilo de tabla

        # Agregar el encabezado de la tabla
        encabezado = tablaCompetenciassinevaluar.rows[0].cells
        encabezado[0].text = 'Competencia'
        encabezado[1].text = 'Resultado de Aprendizaje sin Evaluar'

        # Iterar sobre las filas de la tabla y los datos del DataFrame
        for i, (index, fila) in enumerate(dfcompetenciassinevaluar.iterrows(), start=1):
            # Acceder a las celdas de la fila
            celdas = tablaCompetenciassinevaluar.rows[i].cells
            # Asignar los datos a las celdas
            celdas[0].text = str(fila['Competencia'])
            celdas[1].text = str(fila['Resultado de Aprendizaje'])

    # Insertar la tabla debajo de la frase encontrada
        paragraph._p.addnext(tablaCompetenciassinevaluar._element)
        break


# Agregar un párrafo en blanco entre las tablas
documento_word.add_paragraph()




horas_planeacion = {
    'Utilizar herramientas informáticas de acuerdo con las necesidades de manejo de información': 48,
    'APLICAR PRÁCTICAS DE PROTECCIÓN AMBIENTAL, SEGURIDAD Y SALUD EN EL TRABAJO DE ACUERDO CON LAS POLÍTICAS ORGANIZACIONALES Y LA NORMATIVIDAD VIGENTE.': 48,
    'Razonar cuantitativamente frente a situaciones susceptibles de ser abordadas de manera matemática en contextos laborales, sociales y personales.': 48,
    'APLICACIÓN DE CONOCIMIENTOS DE LAS CIENCIAS NATURALES DE ACUERDO CON SITUACIONES DEL CONTEXTO PRODUCTIVO Y SOCIAL.': 48,
    'DESARROLLAR PROCESOS DE COMUNICACIÓN EFICACES Y EFECTIVOS, TENIENDO EN CUENTA SITUACIONES  DE ORDEN SOCIAL, PERSONAL Y PRODUCTIVO.': 48,
    'GENERAR HÁBITOS SALUDABLES DE VIDA MEDIANTE LA APLICACIÓN DE PROGRAMAS DE ACTIVIDAD FÍSICA EN LOS CONTEXTOS PRODUCTIVOS Y SOCIALES.': 48,
    'Orientar investigación formativa según referentes técnicos': 48,
    'Enrique Low Murtra-Interactuar en el contexto productivo y social de acuerdo con principios  éticos para la construcción de una cultura de paz.': 48,
    'INTERACTUAR EN LENGUA INGLESA DE FORMA ORAL Y ESCRITA DENTRO DE CONTEXTOS SOCIALES Y LABORALES SEGÚN LOS CRITERIOS ESTABLECIDOS POR EL MARCO COMÚN EUROPEO DE REFERENCIA PARA LAS LENGUAS.': 384 ,
    'Fomentar cultura emprendedora según habilidades y competencias personales': 48
}




# Cargar el archivo de Excel en competencias
dfinstructorhoras = pd.read_excel('Reporte de Instructores por Ficha.xls',header=0,skiprows=10)


dfinstructorhoras['Nombrecompleto'] = dfinstructorhoras['Nombre Instructor'] + ' ' + dfinstructorhoras['Apellido Instructor']
dfinstructorhoras = dfinstructorhoras.sort_values(by='Competencia', ascending=True)

# Añadir la columna 'Horas Planeadas' al DataFrame
dfinstructorhoras['Horas Planeadas'] = dfinstructorhoras['Competencia'].map(horas_planeacion)

dfinstructorhoras = dfinstructorhoras[['Nombrecompleto','Estado Instructor','Competencia','Horas Programadas','Horas Planeadas']]



# Buscar la frase dentro de la celda
for paragraph in target_cell.paragraphs:
    if "El reporte de instructor por ficha en SOFIA PLUS es:" in paragraph.text:
            # Agregar una tabla al final del documento
        tablainstructores = documento_word.add_table(rows=len(dfinstructorhoras) + 1, cols=5)  # +1 para incluir el encabezado
        tablainstructores.style = 'Table Grid'  # Aplicar un estilo de tabla

        # Agregar el encabezado de la tabla
        encabezado = tablainstructores.rows[0].cells
        encabezado[0].text = 'Nombre completo Instructor'
        encabezado[1].text = 'Estado Instructor'
        encabezado[2].text = 'Competencia'
        encabezado[3].text = 'Horas Programadas en sofia plus'
        encabezado[4].text = 'Horas Planeacion'

        # Iterar sobre las filas de la tabla y los datos del DataFrame
        for i, (index, fila) in enumerate(dfinstructorhoras.iterrows(), start=1):
            # Acceder a las celdas de la fila
            celdas = tablainstructores.rows[i].cells
            # Asignar los datos a las celdas
            celdas[0].text = str(fila['Nombrecompleto'])
            celdas[1].text = str(fila['Estado Instructor'])
            celdas[2].text = str(fila['Competencia'])
            celdas[3].text = str(fila['Horas Programadas'])
            celdas[4].text = str(fila['Horas Planeadas'])

    # Insertar la tabla debajo de la frase encontrada
        paragraph._p.addnext(tablainstructores._element)
        break


# Agregar un párrafo en blanco entre las tablas
documento_word.add_paragraph()


documento_word.save('acta.docx')

print("Se ha agregado una tabla al final del documento de Word con los datos de las columnas especificadas.")